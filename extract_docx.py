import docx
import os

def extract_docx_content(file_path):
    """
    Extract text content from a .docx file
    """
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

if __name__ == "__main__":
    docx_path = "/home/ubuntu/upload/FreelanceFlow AI.docx"
    output_path = "/home/ubuntu/freelanceflow/requirements.txt"
    
    content = extract_docx_content(docx_path)
    
    # Save the extracted content to a text file
    with open(output_path, "w") as f:
        f.write(content)
    
    print(f"Content extracted and saved to {output_path}")
    print("\nDocument Content Preview:")
    print("="*50)
    print(content[:500] + "..." if len(content) > 500 else content)
    print("="*50)
